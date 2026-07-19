from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "补充材料"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "奇遇长隆_网页逻辑与内容说明_补充材料.docx"
ARCH_IMG = OUT_DIR / "奇遇长隆_系统逻辑图.png"

FOREST = "07382D"
FOREST_DARK = "03271F"
GOLD = "DBAD4C"
MINT = "84D3AC"
CORAL = "E77B63"
CREAM = "F6F1E6"
WARM = "EADFC9"
INK = "13342C"
MUTED = "66766F"
LIGHT = "F4F6F3"
WHITE = "FFFDF8"
GRID = "D8DED9"

# narrative_proposal preset, with two named overrides:
# 1) qiyu_brand_palette: forest/gold/cream brand colors.
# 2) cjk_typeface: Microsoft YaHei for predictable Simplified Chinese rendering.
BASE_FONT = "Microsoft YaHei"
SERIF_FONT = "SimSun"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=GRID, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
    set_table_borders(table)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, font=BASE_FONT) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        p_bdr.append(left)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_image_alt(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description[:80])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def add_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids or [0]) + 1
    next_num = max(num_ids or [0]) + 1

    def build_abstract(abs_id: int, fmt: str, text: str, font: str | None = None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def build_num(num_id: int, abs_id: int):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)

    build_abstract(next_abs, "bullet", "●", BASE_FONT)
    build_num(next_num, next_abs)
    bullet_num = next_num
    build_abstract(next_abs + 1, "decimal", "%1.")
    build_num(next_num + 1, next_abs + 1)
    return bullet_num, next_num + 1


def set_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_bullet(doc: Document, text: str, bullet_num: int, *, lead: str | None = None):
    p = doc.add_paragraph()
    set_num(p, bullet_num)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    if lead and text.startswith(lead):
        r1 = p.add_run(lead)
        set_run_font(r1, size=10.5, color=INK, bold=True)
        r2 = p.add_run(text[len(lead):])
        set_run_font(r2, size=10.5, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=INK)
    return p


def add_step(doc: Document, text: str, decimal_num: int):
    p = doc.add_paragraph()
    set_num(p, decimal_num)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    keep_with_next(p)
    return p


def add_kicker(doc: Document, text: str, after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text.upper())
    set_run_font(r, size=9.5, color=GOLD, bold=True)
    r.font.all_caps = True
    return p


def add_callout(doc: Document, label: str, text: str, *, fill=LIGHT, border=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, fill, border)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.5, color=FOREST, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    keep_with_next(p)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)
    return p


def add_figure(doc: Document, path: Path, caption: str, alt: str, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_image_alt(shape, alt)
    add_caption(doc, caption)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], FOREST)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[idx], "F7F8F6")
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(value)
            set_run_font(r, size=9, color=INK, bold=(idx == 0))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_section_page(doc: Document, kicker: str, title: str, lead: str | None = None):
    doc.add_page_break()
    add_kicker(doc, kicker)
    add_heading(doc, title, 1)
    if lead:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(lead)
        set_run_font(r, size=12.5, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    tokens = {
        "Heading 1": (16, FOREST, 18, 10),
        "Heading 2": (13, FOREST, 12, 6),
        "Heading 3": (12, INK, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r1 = hp.add_run("奇遇长隆  |  网页逻辑与内容说明")
    set_run_font(r1, size=8.5, color=MUTED, bold=True)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def create_architecture_diagram() -> None:
    w, h = 1800, 860
    img = Image.new("RGB", (w, h), "#F6F1E6")
    draw = ImageDraw.Draw(img)
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    bold_path = Path("C:/Windows/Fonts/msyhbd.ttc")
    title_font = ImageFont.truetype(str(bold_path), 52)
    box_title = ImageFont.truetype(str(bold_path), 30)
    body_font = ImageFont.truetype(str(font_path), 24)
    small_font = ImageFont.truetype(str(font_path), 21)

    draw.text((90, 55), "奇遇长隆：从用户输入到业务动作的统一 Agent 闭环", font=title_font, fill="#07382D")
    draw.rounded_rectangle((90, 150, 1710, 250), radius=28, fill="#07382D")
    state_text = "共享旅程状态：游客画像  ·  伙伴身份  ·  动态路线  ·  任务成就  ·  授权媒体  ·  可解释决策记录"
    bbox = draw.textbbox((0, 0), state_text, font=body_font)
    draw.text(((w - (bbox[2] - bbox[0])) / 2, 184), state_text, font=body_font, fill="#FFFDF8")

    boxes = [
        (90, 330, 390, 590, "用户与环境输入", ["同行成员 / 节奏偏好", "位置 / 客流 / 天气", "现场反馈 / 照片授权"]),
        (465, 330, 765, 590, "全局行程 Agent", ["多约束路线生成", "实时重排与原因解释", "休息 / 用餐 / 服务插入"]),
        (840, 330, 1140, 590, "空间动物 Agent", ["50m 电子围栏触发", "角色化科普与任务", "年龄分层 / 内容约束"]),
        (1215, 330, 1515, 590, "记忆生成引擎", ["足迹与徽章聚合", "AIGC 回忆短片", "个性票根与分享文案"]),
    ]
    fills = ["#EADFC9", "#D8E9DF", "#E9E1C9", "#F0DCD6"]
    for i, (x1, y1, x2, y2, title, lines) in enumerate(boxes):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=28, fill=fills[i], outline="#07382D", width=3)
        draw.text((x1 + 28, y1 + 30), title, font=box_title, fill="#07382D")
        for j, line in enumerate(lines):
            draw.ellipse((x1 + 30, y1 + 105 + j * 47, x1 + 43, y1 + 118 + j * 47), fill="#DBAD4C")
            draw.text((x1 + 57, y1 + 96 + j * 47), line, font=small_font, fill="#13342C")
        if i < len(boxes) - 1:
            ax1, ay = x2 + 13, (y1 + y2) // 2
            ax2 = boxes[i + 1][0] - 13
            draw.line((ax1, ay, ax2, ay), fill="#DBAD4C", width=7)
            draw.polygon([(ax2, ay), (ax2 - 20, ay - 13), (ax2 - 20, ay + 13)], fill="#DBAD4C")

    draw.rounded_rectangle((320, 680, 1480, 785), radius=28, fill="#FFFDF8", outline="#DBAD4C", width=3)
    output = "页面反馈与业务动作：路线更新 · 主动讲解 · 服务提醒 · 餐饮/零售承接 · 游后分享 · 复游激励"
    obox = draw.textbbox((0, 0), output, font=body_font)
    draw.text(((w - (obox[2] - obox[0])) / 2, 717), output, font=body_font, fill="#07382D")
    for x in (240, 615, 990, 1365):
        draw.line((x, 592, x, 670), fill="#84D3AC", width=6)
        draw.polygon([(x, 675), (x - 13, 653), (x + 13, 653)], fill="#84D3AC")
    img.save(ARCH_IMG, quality=95)


def build_document() -> None:
    create_architecture_diagram()
    doc = Document()
    configure_document(doc)
    configure_styles(doc)
    bullet_num, decimal_num = add_numbering(doc)

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(46)
    add_kicker(doc, "2026 AI 先锋未来人才大赛｜长隆集团命题", after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("奇遇长隆")
    set_run_font(r, size=32, color=FOREST, bold=True, font=SERIF_FONT)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(20)
    r = p2.add_run("全周期 AI 动物伙伴网页逻辑与内容说明")
    set_run_font(r, size=17, color=INK, bold=True)

    panda = ROOT / "qiyu-remotion-video" / "public" / "images" / "tuantuan.png"
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = pic_p.add_run().add_picture(str(panda), width=Inches(2.25))
    set_image_alt(shape, "奇遇长隆熊猫 AI 伙伴团团")

    slogan = doc.add_paragraph()
    slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    slogan.paragraph_format.space_before = Pt(8)
    slogan.paragraph_format.space_after = Pt(28)
    sr = slogan.add_run("“让每一次出发，都成为一场奇遇。”")
    set_run_font(sr, size=14, color=FOREST, italic=True, font=SERIF_FONT)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.line_spacing = 1.5
    for idx, line in enumerate([
        "项目形态：移动端 H5 / 高保真交互原型",
        "核心客群：亲子家庭",
        "演示链路：游前规划 → 游中陪伴 → 游后留存",
        "版本日期：2026 年 7 月",
    ]):
        rr = meta.add_run(line + ("\n" if idx < 3 else ""))
        set_run_font(rr, size=10.5, color=MUTED, bold=(idx == 0))

    # 1. Executive summary
    add_section_page(doc, "01 / 项目总览", "从“被动导览”到“持续陪伴”", "网页展示的是一条可被实际体验、可被业务验证的全周期 Agent 链路。")
    add_callout(doc, "一句话定义", "为每组游客匹配一位可持续记忆的 AI 动物伙伴：游前共创路线，游中感知位置与客流并动态改线，游后将足迹、互动与照片生成可珍藏、可分享的旅行记忆。")
    add_body(doc, "长隆命题要求从“游前感知—游中沉浸—游后留存”的全周期游客视角出发，用 AI 重塑人与乐园空间的连接，并让体验自然穿梭于酒店、乐园、餐饮、零售与娱乐衍生品。奇遇长隆以“统一旅程状态 + 多 Agent 协作 + 空间触发 + AIGC 记忆生成”回应这一要求。")
    add_heading(doc, "核心问题", 2)
    for text, lead in [
        ("信息分散：地图、排队、餐饮、演出与任务彼此割裂，游客仍要自己完成复杂决策。", "信息分散："),
        ("推荐静态：传统路线无法响应疲劳、拥堵、临时如厕、天气变化等现场变量。", "推荐静态："),
        ("互动被动：动物讲解往往是固定内容，缺少位置、角色、年龄与连续记忆。", "互动被动："),
        ("旅程中断：离园后照片与轨迹散落，情绪价值无法沉淀为传播、复游和消费。", "旅程中断："),
    ]:
        add_bullet(doc, text, bullet_num, lead=lead)
    add_heading(doc, "三阶段价值闭环", 2)
    add_table(
        doc,
        ["阶段", "用户输入", "Agent 行动", "页面输出与价值"],
        [
            ["游前", "同行结构、节奏、时间、动物与饮食偏好", "需求访谈、路线计算、推荐解释", "专属路线与伙伴绑定；降低决策负担"],
            ["游中", "位置、排队、任务进度、现场反馈", "地理围栏触发、科普互动、动态重排", "实时地图、区域任务与改线；提升沉浸和效率"],
            ["游后", "足迹、徽章、伙伴留言、授权照片", "内容聚合、AIGC 叙事、票根生成", "回忆短片、个性票根与分享；促进传播和复游"],
        ],
        [1100, 2450, 2600, 3210],
    )
    add_body(doc, "命题来源：2026 AI 先锋未来人才大赛“长隆集团”命题页（访问日期：2026-07-19）。网页原型以竞赛场景演示为目的，涉及客流、排队、定位和商业接口的数据均在文末明确区分“仿真验证”与“生产接入”。", after=0)

    # 2. Unified logic
    add_section_page(doc, "02 / 系统逻辑", "一份旅程状态，驱动四类 Agent 能力", "核心不在页面数量，而在数据与记忆能否跨页面连续。")
    add_figure(doc, ARCH_IMG, "图 1  奇遇长隆统一 Agent 闭环", "用户输入、全局行程 Agent、空间动物 Agent、记忆生成引擎与业务动作之间的闭环逻辑", 6.5)
    add_heading(doc, "统一旅程状态", 2)
    for text, lead in [
        ("journeyProfile｜游客画像：同行成员、年龄、节奏、时间窗、兴趣、饮食及特殊需求。", "journeyProfile｜游客画像："),
        ("companion｜伙伴身份：所选动物 Agent、性格语气、关系进度与历史记忆。", "companion｜伙伴身份："),
        ("itinerary｜动态行程：站点顺序、预计时间、路网距离、推荐理由与备选点。", "itinerary｜动态行程："),
        ("liveContext｜实时环境：当前位置、地理围栏、客流、排队、天气和现场反馈。", "liveContext｜实时环境："),
        ("achievements / media｜成果资产：任务、徽章、照片授权、伙伴留言、回忆短片与票根。", "achievements / media｜成果资产："),
    ]:
        add_bullet(doc, text, bullet_num, lead=lead)
    add_callout(doc, "连续性的关键", "游前不是一次问卷，游中不是另一张地图，游后也不是孤立相册。每一次选择、抵达、提问和任务完成都会写回同一份 journeyRecord，并成为下一阶段的上下文。", fill="F7F2E5", border=GOLD)
    add_heading(doc, "三条产品原则", 2)
    add_body(doc, "1）先理解再推荐：每一条路线都说明“为什么适合你”；2）主动但可控：Agent 可提出调整，但保留用户确认与撤销；3）把位置变成内容入口：只有抵达真实展区，才唤醒对应的动物故事、任务和知识。")

    # 3. Information architecture
    add_section_page(doc, "03 / 页面架构", "四个核心路由，完成一次完整旅程", "网页以移动端纵向体验为主，通过明确的阶段路由和共享状态串联。")
    add_table(
        doc,
        ["路由", "用户目标", "关键触发", "核心输出"],
        [
            ["/pretrip", "在入园前快速形成可执行计划", "完成伙伴与 7 轮需求采集后点击“生成我的路线”", "伙伴绑定、路线地图、时间线、站点理由与备选方案"],
            ["/inpark", "按计划探索，并随现场变化调整", "选择“按计划探险”；接近展区进入 50m 围栏", "实时地图、抵达确认、区域任务、AI 科普、动态改线"],
            ["/posttrip", "回看今日足迹并沉淀成果", "行程结束或从底部导航进入“结束”", "计划/实际对照、徽章、照片位、AIGC 回忆短片入口"],
            ["票根编辑页", "把旅程转成可保存、可分享的纪念物", "点击“生成奇遇票根”", "三种票根样式、封面、标题留言、分享文案与导出"],
        ],
        [1350, 2500, 2760, 2750],
    )
    add_heading(doc, "全局导航与状态流转", 2)
    add_body(doc, "底部导航保留“地图—服务—伙伴—结束”四个稳定入口，既保证主线任务，又允许游客在真实园区中随时偏离路线。核心状态依次为：draft（采集中）→ ready（路线已生成）→ active（游园中）→ adjusted（发生过改线）→ completed（已结束）→ archived（已生成记忆资产）。")
    add_heading(doc, "页面共通反馈", 2)
    for text in [
        "进度可见：游前显示 1/7 等对话轮次，游中显示 0/3 站与伙伴任务进度。",
        "原因可见：路线站点与改线结果均展示距离、等待、适配原因或被替换原因。",
        "状态可逆：允许重选伙伴、回到实时地图、重新生成路线/票根，减少一步走错的焦虑。",
        "异常有兜底：无定位、知识不确定、紧急服务或未授权照片时，页面提供替代入口而非中断体验。",
    ]:
        add_bullet(doc, text, bullet_num)

    # 4. Pretrip
    add_section_page(doc, "04 / 游前逻辑", "AI 逐问逐答，把模糊需求变成真实路线", "团团不是先给答案，而是通过短对话逐步建立可计算的旅程约束。")
    pretrip_img = ROOT / "qiyu-remotion-video" / "output" / "stills" / "01-pretrip.png"
    add_figure(doc, pretrip_img, "图 2  游前：通过自然对话理解同行结构与偏好", "奇遇长隆游前页面，展示团团与亲子家庭进行需求访谈", 6.5)
    add_heading(doc, "对话采集顺序", 2)
    steps = [
        "选择动物伙伴：团团（温柔亲子科普）、凯凯（行动型探险）、悠米（松弛休闲）。",
        "确认同行结构：亲子家庭、情侣、朋友结伴或独自出游，并记录成人/儿童人数。",
        "选择游览节奏：效率优先、适度探索或悠享慢游。",
        "设置时间窗：预计入园与离园时间，形成路线硬约束。",
        "按优先级选择动物：本次演示为熊猫园、考拉馆、长颈鹿园。",
        "确认园内用餐需求：为后续插入用餐节点与错峰建议提供依据。",
        "补充自由文本：如“孩子下午容易累，希望少走回头路”。",
        "确认摘要并生成路线：Agent 汇总条件，用户可返回修改后再计算。",
    ]
    for step in steps:
        add_step(doc, step, decimal_num)

    add_heading(doc, "路线生成逻辑", 2)
    add_callout(doc, "概念评分", "站点评分 = 偏好匹配 + 动物活跃时段 + 时间窗适配 + 科普/互动价值 − 步行成本 − 排队风险 − 疲劳惩罚。随后在园区路网中计算满足硬约束的低回头路线，并保留弹性点与备选点。")
    add_body(doc, "输出不只是一张地图，还包括：总里程与预计时长、时间轴、每段步行距离、每一站的推荐理由、用餐/休息节点以及遇到高峰或闭园时的替代项。当前原型用模拟站点、路网与排队数据验证交互闭环；生产版需接入真实 GIS、营业时间、动物活跃时段、演出和餐饮接口。")

    # 5. In-park
    add_section_page(doc, "05 / 游中逻辑", "抵达的那一刻，知识主动发生", "全局行程 Agent 负责调度；空间动物 Agent 只在正确的地点、正确的时机被唤醒。")
    inpark_img = ROOT / "qiyu-remotion-video" / "output" / "stills" / "02-inpark.png"
    add_figure(doc, inpark_img, "图 3  游中：实时路线、50m 电子围栏与区域互动", "奇遇长隆游中页面，展示实时地图、路线和地理围栏", 6.5)
    add_heading(doc, "演示链路与状态变化", 2)
    for step in [
        "选择“按计划探险”后加载游前路线，进入 LIVE GIS 实时地图。",
        "点击“模拟前往长颈鹿园”，位置沿路线变化；进入 50m 围栏后出现“演示抵达”。",
        "确认抵达后，系统解锁“开始这一站”，并写入实际足迹与到达时间。",
        "进入长颈鹿园现场奇遇站，完成“高处取食挑战”，选择“舌头”获得任务反馈与徽章进度。",
        "向“高空瞭望员”提问“长颈鹿的舌头为什么这么长？”，由受约束的知识问答返回答案。",
        "返回实时地图，选择“前面排队很长，帮我调整路线”，全局 Agent 重排后续顺序并解释变化。",
    ]:
        add_step(doc, step, decimal_num)
    add_heading(doc, "两类触发机制", 2)
    add_body(doc, "空间触发：电子围栏（原型为 50m）负责判断“游客是否真正抵达”，二维码/NFC 可作为定位不稳定时的备用入口。事件触发：客流升高、天气变化、游客疲劳、临时如厕或项目关闭等事件写入 liveContext，触发路线重新评估。")
    add_heading(doc, "Agent 回答边界", 2)
    for text in [
        "动物知识来自经过审核的知识库，回答必须可追溯，并明确区分事实、推测与趣味表达。",
        "根据儿童/成人切换表达深度；不提供危险接近、投喂或干扰动物的建议。",
        "对医疗、安全、走失等紧急情境直接提供人工服务入口，不由生成式模型独立处置。",
        "改线建议必须说明原因与影响，用户可保持原计划或撤销调整。",
    ]:
        add_bullet(doc, text, bullet_num)

    # 6. Posttrip
    add_section_page(doc, "06 / 游后逻辑", "把零散瞬间，编织成完整旅行故事", "离园不是服务终点，而是记忆沉淀、社交传播和复游运营的起点。")
    posttrip_img = ROOT / "qiyu-remotion-video" / "output" / "stills" / "03-posttrip.png"
    add_figure(doc, posttrip_img, "图 4  游后：足迹、成就、回忆短片与票根入口", "奇遇长隆游后页面，展示照片位、回忆短片和奇遇票根", 6.5)
    add_heading(doc, "内容聚合逻辑", 2)
    for text, lead in [
        ("计划与实际：对照原始路线和真实足迹，保留发生过的改线与临时发现。", "计划与实际："),
        ("成就与证据：把展区抵达、观察任务、问答和徽章组成可回看的“认真观察过的证明”。", "成就与证据："),
        ("伙伴叙事：动物 Agent 根据当天互动生成短留言，使回忆保持角色一致性。", "伙伴叙事："),
        ("授权媒体：仅使用游客主动选择的照片/视频；没有照片时仍可用地图、徽章和插画完成基础版。", "授权媒体："),
    ]:
        add_bullet(doc, text, bullet_num, lead=lead)
    add_heading(doc, "AIGC 回忆短片", 2)
    add_body(doc, "记忆引擎按照“出发—抵达—互动—变化—完成”的叙事节奏，把路线足迹、徽章、伙伴留言、票根和授权照片组织为 15—30 秒的 H5 回忆短片。用户可预览、逐幕查看、替换素材或关闭，不把“自动生成”变成不可控的黑箱。")
    add_heading(doc, "奇遇票根", 2)
    add_body(doc, "票根编辑页提供三种风格，包括动物伙伴票。用户可设置封面照片、缩放裁切、标题、旅程留言与分享文案，并导出长图或 9:16 视频封面。票根将真实行程数据压缩为一个可保存、可分享、可制作实体纪念品的情绪载体。")

    # 7. Content system
    add_section_page(doc, "07 / 内容系统", "同一套园区知识，不同的角色与陪伴策略", "多 Agent 的差异不只是头像，而是语气、路线策略、任务类型和情绪反馈的系统性差异。")
    add_table(
        doc,
        ["伙伴", "角色定位", "路线策略", "内容与适用人群"],
        [
            ["团团｜熊猫", "温柔耐心的亲子科普官", "节奏舒缓、休息充足、少回头", "短句、类比、亲子问答；适合儿童家庭"],
            ["凯凯｜白虎", "行动果断的路线探险队长", "效率优先、重点明确、挑战导向", "任务挑战、探索成就；适合青年与朋友结伴"],
            ["悠米｜考拉", "松弛细心的休闲路线管家", "低强度、拍照与停留友好", "情绪陪伴、慢游提醒；适合情侣、老人或轻松游"],
        ],
        [1250, 2100, 2500, 3510],
    )
    add_heading(doc, "内容分层", 2)
    for text, lead in [
        ("基础服务层：路线、距离、等待、营业时间、厕所、母婴室、餐饮和演出提醒。", "基础服务层："),
        ("沉浸互动层：角色问候、在地科普、观察任务、即时问答、徽章与好感度。", "沉浸互动层："),
        ("情绪记忆层：伙伴寄语、回忆短片、票根、动物图鉴与下一次相见的伏笔。", "情绪记忆层："),
        ("商业承接层：与已发生的需求和情绪相关的餐饮、零售、酒店与纪念品建议。", "商业承接层："),
    ]:
        add_bullet(doc, text, bullet_num, lead=lead)
    add_heading(doc, "内容治理", 2)
    add_body(doc, "知识内容采用“官方资料/专家审核知识库 + 检索增强生成 + 规则约束”的组合；角色可以有温度，但事实不能被人设改写。系统记录知识来源、版本与敏感规则，对不确定问题明确说明并提供人工咨询入口。未成年人、定位、照片和生物安全相关内容采用更严格的默认权限与提示。")
    add_callout(doc, "文案原则", "伙伴要像一个真正同行的人：主动但不打扰、温柔但不幼稚、解释但不说教、推荐但不硬卖。")

    # 8. Business loop
    add_section_page(doc, "08 / 多业态闭环", "让情绪价值自然转化为服务与消费", "商业转化不是额外弹窗，而是把正确的服务放到正确的时间、地点和情境中。")
    add_table(
        doc,
        ["业态", "真实需求触发", "Agent 服务", "价值与边界"],
        [
            ["酒店", "入园前准备、体力恢复、次日衔接", "房型/套餐建议、班车与入住指引、睡前故事、次日路线", "延长旅程周期；需接入库存与订单，保持用户自主选择"],
            ["餐饮", "临近用餐、儿童/过敏需求、附近拥堵", "错峰推荐、替代餐厅、排号/点单提醒", "节省等待并提升转化；不以优惠轰炸替代服务"],
            ["零售", "完成动物互动或获得徽章", "推荐对应 IP 周边、收藏、自提或快递", "商品承接当日情绪；明确库存、价格与授权"],
            ["娱乐", "演出临近、路线空档、天气变化", "场次衔接、步行时间、替代室内体验", "提高内容触达和停留；避免过度排程"],
            ["纪念品", "行程完成、票根/短片生成", "电子内容升级、定制画册、明信片或徽章", "把数字记忆实体化；生成与购买分步确认"],
        ],
        [1100, 2460, 2800, 3000],
    )
    add_heading(doc, "商业闭环的三次转化", 2)
    add_body(doc, "第一次是服务转化：用错峰、排号、导航解决即时问题；第二次是情绪转化：把刚完成的互动变成主题商品或定制纪念；第三次是关系转化：用动物伙伴记忆、未完成图鉴和季节新路线推动会员与复游。每一次转化均由用户需求或旅程事件触发，并提供清晰的跳过、撤回与退订。")
    add_callout(doc, "核心判断", "如果推荐不能同时提升用户体验，就不应被视为合格的商业触点。", fill="F7F2E5", border=GOLD)

    # 9. Technical architecture
    add_section_page(doc, "09 / 数据与技术", "原型验证闭环，生产环境连接真实业务", "当前网页证明交互与系统逻辑可行；规模化落地依赖数据接口、内容治理和隐私机制。")
    add_heading(doc, "能力分层", 2)
    for text, lead in [
        ("体验层：移动端 H5，包括对话式采集、地图路线、展区奇遇站、游后回顾与票根编辑。", "体验层："),
        ("Agent 编排层：统一管理伙伴人设、旅程记忆、工具调用、事件触发、权限和人工兜底。", "Agent 编排层："),
        ("智能能力层：多约束路线规划、地理围栏、动物知识 RAG、内容审核、AIGC 图像/视频/文案。", "智能能力层："),
        ("数据服务层：GIS 路网、项目与展区、客流排队、天气演出、餐饮零售、酒店票务和会员。", "数据服务层："),
        ("治理层：定位与媒体授权、未成年人保护、数据最小化、日志审计、知识版本和模型安全。", "治理层："),
    ]:
        add_bullet(doc, text, bullet_num, lead=lead)
    add_heading(doc, "原型数据与生产数据边界", 2)
    add_table(
        doc,
        ["能力", "当前原型", "生产落地要求"],
        [
            ["路线与地图", "模拟路网、站点距离和示例路线", "真实 GIS 路网、无障碍路径、开放时间与接驳信息"],
            ["客流与改线", "按钮触发的拥堵事件和仿真等待时间", "实时/预测客流、排队接口、演出与天气事件流"],
            ["空间触发", "“模拟前往/演示抵达”验证 50m 围栏流程", "用户授权定位、蓝牙信标或二维码/NFC 备用确认"],
            ["动物问答", "受约束示例回答与角色化输出", "官方知识库、专家审核、来源追溯、敏感规则与人工兜底"],
            ["AIGC 内容", "预设旅程数据、样例照片与票根模板", "素材授权、生成审核、编辑删除、导出与存储生命周期"],
            ["商业服务", "展示推荐与跳转逻辑", "库存、排号、订单、支付、会员和售后等企业 API"],
        ],
        [1700, 3400, 4260],
    )

    # 10. Evaluation and roadmap
    add_section_page(doc, "10 / 验证指标与路线图", "用体验、运营与商业三组指标证明价值", "原型的目标不是证明所有接口已经存在，而是证明一个值得接入真实数据的业务闭环。")
    add_heading(doc, "建议指标", 2)
    add_table(
        doc,
        ["维度", "核心指标", "验证问题"],
        [
            ["体验效率", "路线生成完成率、平均步行/排队时间、改线接受率、任务中断率", "是否减少决策负担，并让变化更可控？"],
            ["沉浸互动", "围栏唤醒率、区域任务完成率、问答继续率、徽章完成率", "位置与角色是否真正提升观察和科普兴趣？"],
            ["内容传播", "短片生成率、编辑率、保存/分享率、票根二次打开率", "游后内容是否值得保存并愿意主动传播？"],
            ["商业价值", "餐饮/零售点击与转化、纪念品升级率、会员绑定与复游意愿", "推荐是否同时提升体验与客单价？"],
            ["运营健康", "知识纠错率、人工转接率、授权撤回率、投诉与敏感问题命中率", "系统是否可信、可治理、可长期运营？"],
        ],
        [1600, 3920, 3840],
    )
    add_heading(doc, "分阶段路线图", 2)
    for step in [
        "MVP 验证：完善三段网页闭环，以仿真数据完成 5—8 组亲子家庭可用性测试。",
        "小范围试点：选择一个园区与 2—3 个明星动物展区，接入地图、营业时间和有限实时数据。",
        "业务联动：接入餐饮排号、零售库存、票根纪念品和酒店/班车信息，验证二次消费指标。",
        "规模化运营：建立内容审核后台、跨园区数字身份、会员权益和长期动物伙伴记忆。",
    ]:
        add_step(doc, step, decimal_num)
    add_heading(doc, "主要风险与应对", 2)
    add_body(doc, "定位漂移通过“围栏 + 信标/扫码确认”降低误触发；路线建议通过原因说明和用户确认降低自动化焦虑；动物知识以官方/专家审核与来源追溯控制事实风险；未成年人、照片与位置数据采用最小必要采集、分级授权、可删除与限期存储；涉及支付、医疗、安全和紧急事件时由业务系统或人工服务兜底。")

    # 11. Demo guide
    add_section_page(doc, "11 / 演示说明", "三分钟讲清“感知—沉浸—留存”", "评审看到的不只是三张页面，而是一份连续增长、持续记忆的 journeyRecord。")
    add_heading(doc, "推荐演示脚本", 2)
    demo_steps = [
        "游前（60 秒）：选择团团与亲子家庭，设为悠享慢游、10:00—18:00，选择熊猫/考拉/长颈鹿，勾选园内用餐并补充“孩子下午容易累，希望少走回头路”，生成路线。",
        "游中（80 秒）：选择按计划探险，展示实时地图与 50m 围栏；模拟抵达长颈鹿园，完成高处取食挑战并提问；返回地图，报告排队很长，展示改线结果。",
        "游后（40 秒）：展示计划与实际、成就徽章和照片位；播放回忆短片，进入票根编辑页，切换“动物伙伴票”，展示分享文案与导出入口。",
    ]
    for step in demo_steps:
        add_step(doc, step, decimal_num)
    add_heading(doc, "演示时应强调的四句话", 2)
    for text in [
        "我们不是在地图旁边增加一个 AI 对话框，而是让同一个 Agent 贯穿整个旅程。",
        "位置不是一个坐标，而是动物知识、任务与角色故事的入口。",
        "每一次改线都可解释、可选择；每一次内容生成都建立在用户授权之上。",
        "商业推荐由真实需求和情绪节点触发，既提升体验，也为多业态创造自然增量。",
    ]:
        add_bullet(doc, text, bullet_num)
    add_callout(doc, "结语", "奇遇长隆把“规划工具”升级为“关系型服务”：它记得游客是谁、此刻在哪里、今天发生了什么，也知道下一次为什么值得再来。", fill=FOREST, border=GOLD)
    # Change callout text to white on dark fill.
    last_p = doc.paragraphs[-1]
    for idx, run in enumerate(last_p.runs):
        set_run_font(run, size=10.5, color=(GOLD if idx == 0 else WHITE), bold=(idx == 0))

    add_heading(doc, "原型与素材索引", 2)
    add_body(doc, "网页演示地址：qiyucl.site/pretrip、qiyucl.site/inpark、qiyucl.site/posttrip。配套宣传片、页面录屏、角色图与票根素材已用于本说明书插图。比赛命题依据：https://activity.feishu.cn/future-talent?detail=chimelonggroup。", after=0)

    # Core properties and final save.
    doc.core_properties.title = "奇遇长隆｜全周期AI动物伙伴网页逻辑与内容说明"
    doc.core_properties.subject = "2026 AI先锋未来人才大赛长隆集团命题补充材料"
    doc.core_properties.author = "奇遇长隆项目团队"
    doc.core_properties.keywords = "长隆, AI Agent, 智慧伴游, 游前游中游后, AIGC"
    doc.core_properties.comments = "参赛补充材料：网页逻辑、内容体系、技术边界与业务价值说明"
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
    print(DOCX_PATH)
